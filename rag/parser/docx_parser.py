"""
python-docx를 사용한 DOCX 문서 텍스트 추출 모듈.

DOCX 파일의 본문 문단과 테이블 셀에서 텍스트를 추출합니다.
빈 문단은 건너뛰고, 문단 간 줄바꿈으로 연결합니다.
"""

import logging
from pathlib import Path

from rag.parser import EncryptedFileError

logger = logging.getLogger(__name__)


def _is_encrypted_ooxml(file_path: str) -> bool:
    """OOXML 파일이 MS-OFFCRYPTO 암호화 상태인지 확인 (OLE2 + EncryptedPackage)."""
    try:
        import olefile
        if olefile.isOleFile(file_path):
            ole = olefile.OleFileIO(file_path)
            encrypted = ole.exists("EncryptedPackage")
            ole.close()
            return encrypted
    except Exception:
        pass
    return False


def parse_docx(file_path: str) -> str:
    """
    DOCX 파일에서 본문 문단 + 테이블 텍스트를 추출하여 하나의 문자열로 결합.

    python-docx의 Document 객체를 통해 본문 문단(paragraphs)과
    테이블(tables) 내 셀 텍스트를 순서대로 추출합니다.
    빈 문단/셀은 건너뛰고, 줄바꿈(\\n)으로 연결합니다.

    Args:
        file_path: DOCX 파일 절대 경로.

    Returns:
        추출된 전체 텍스트 문자열.

    Raises:
        FileNotFoundError: 파일이 존재하지 않을 때.
        ImportError: python-docx가 설치되지 않았을 때.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX not found: {file_path}")

    # 암호화 감지 (OOXML 암호화 시 OLE2 컨테이너로 저장됨)
    if _is_encrypted_ooxml(file_path):
        raise EncryptedFileError(file_path, "docx")

    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not installed. Cannot parse DOCX files.")
        raise

    try:
        doc = Document(file_path)
        parts = []

        # 본문 문단 추출
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # 테이블 내 셀 텍스트 추출
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    parts.append(" | ".join(row_texts))

        full_text = "\n".join(parts)
        logger.info(f"DOCX parsed: {path.name} ({len(parts)} sections, {len(full_text)} chars)")
        return full_text

    except Exception:
        logger.exception(f"Failed to parse DOCX: {file_path}")
        raise
